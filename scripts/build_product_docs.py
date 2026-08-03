from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SPEC_PATH = ROOT / "ТЗ" / "Спецификация Simple WMS 0.1.docx"
TERMS_PATH = ROOT / "ТЗ" / "Термины и определения.docx"

FONT = "Arial"
ACCENT = RGBColor(31, 78, 121)
ACCENT_DARK = RGBColor(15, 44, 70)
MUTED = RGBColor(90, 100, 112)
LIGHT_FILL = "E8EEF5"
VERY_LIGHT_FILL = "F4F6F9"
WHITE = RGBColor(255, 255, 255)
BLACK = RGBColor(20, 24, 28)

# A4 override for the standard_business_brief preset.
PAGE_WIDTH_DXA = 11906
PAGE_MARGIN_DXA = 1134
CONTENT_WIDTH_DXA = PAGE_WIDTH_DXA - 2 * PAGE_MARGIN_DXA
TABLE_INDENT_DXA = 120


def set_run_font(
    run,
    *,
    size: float | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    color: RGBColor | None = None,
) -> None:
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str = "C9D1D9", size: str = "6") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def set_keep_with_next(paragraph) -> None:
    paragraph.paragraph_format.keep_with_next = True


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    heading_tokens = {
        "Heading 1": (16, ACCENT, 16, 8),
        "Heading 2": (13, ACCENT, 12, 6),
        "Heading 3": (11.5, ACCENT_DARK, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.167

    if "Lead Callout" not in styles:
        callout = styles.add_style("Lead Callout", WD_STYLE_TYPE.PARAGRAPH)
    else:
        callout = styles["Lead Callout"]
    callout.font.name = FONT
    callout._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    callout._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    callout._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    callout.font.size = Pt(11)
    callout.font.bold = True
    callout.font.color.rgb = ACCENT_DARK
    callout.paragraph_format.left_indent = Cm(0.4)
    callout.paragraph_format.right_indent = Cm(0.4)
    callout.paragraph_format.space_before = Pt(6)
    callout.paragraph_format.space_after = Pt(10)
    callout.paragraph_format.line_spacing = 1.15

    if "Table Text" not in styles:
        table_text = styles.add_style("Table Text", WD_STYLE_TYPE.PARAGRAPH)
    else:
        table_text = styles["Table Text"]
    table_text.font.name = FONT
    table_text._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    table_text._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    table_text._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    table_text.font.size = Pt(9.5)
    table_text.paragraph_format.space_before = Pt(0)
    table_text.paragraph_format.space_after = Pt(0)
    table_text.paragraph_format.line_spacing = 1.08


def configure_document(doc: Document, running_label: str) -> None:
    configure_styles(doc)
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)
    section.different_first_page_header_footer = True

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(running_label)
    set_run_font(run, size=8.5, color=MUTED)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("Страница ")
    set_run_font(run, size=8.5, color=MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    p._p.append(field)

    core = doc.core_properties
    core.author = "Simple WMS"
    core.subject = "Продуктовая спецификация автономной складской системы"
    core.keywords = "WMS, склад, адресное хранение, 1С, интеграция"
    core.comments = "Сформировано из спецификации Simple WMS 0.1"


def add_cover(
    doc: Document,
    *,
    kicker: str,
    title: str,
    subtitle: str,
    status: str,
    version: str,
) -> None:
    for _ in range(3):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    run = p.add_run(kicker.upper())
    set_run_font(run, size=10.5, bold=True, color=ACCENT)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(title)
    set_run_font(run, size=28, bold=True, color=ACCENT_DARK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(30)
    run = p.add_run(subtitle)
    set_run_font(run, size=14, color=MUTED)

    meta = [
        ("Статус", status),
        ("Версия продукта", version),
        ("Дата редакции", date.today().strftime("%d.%m.%Y")),
        ("Назначение", "Базовый документ для разработки, приёмки и продуктового развития"),
    ]
    for label, value in meta:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(5)
        r = p.add_run(f"{label}: ")
        set_run_font(r, size=10.5, bold=True, color=BLACK)
        r = p.add_run(value)
        set_run_font(r, size=10.5, color=BLACK)

    doc.add_page_break()


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph(text, style=f"Heading {level}")
    set_keep_with_next(p)


def add_paragraph(doc: Document, text: str, *, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        first, rest = text.split(":", 1)
        r = p.add_run(first + ":")
        set_run_font(r, bold=True)
        r = p.add_run(rest)
        set_run_font(r)
    else:
        r = p.add_run(text)
        set_run_font(r)


def add_callout(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Lead Callout")
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), VERY_LIGHT_FILL)
    p_pr.append(shd)
    r = p.add_run(text)
    set_run_font(r, size=11, bold=True, color=ACCENT_DARK)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_run_font(r)


def create_decimal_numbering(doc: Document) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
    ]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    level.append(num_fmt)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "%1.")
    level.append(level_text)
    suffix = OxmlElement("w:suff")
    suffix.set(qn("w:val"), "tab")
    level.append(suffix)

    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    p_pr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "720")
    indent.set(qn("w:hanging"), "360")
    p_pr.append(indent)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    level.append(p_pr)

    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), FONT)
    fonts.set(qn("w:hAnsi"), FONT)
    fonts.set(qn("w:eastAsia"), FONT)
    r_pr.append(fonts)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "21")
    r_pr.append(size)
    level.append(r_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_numbers(doc: Document, items: list[str]) -> None:
    num_id = create_decimal_numbering(doc)
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.167
        p_pr = p._p.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_id_node = OxmlElement("w:numId")
        num_id_node.set(qn("w:val"), str(num_id))
        num_pr.append(ilvl)
        num_pr.append(num_id_node)
        p_pr.append(num_pr)
        r = p.add_run(item)
        set_run_font(r)


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    widths_dxa: list[int],
) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths_dxa)
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_FILL)
        p = cell.paragraphs[0]
        p.style = doc.styles["Table Text"]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(header)
        set_run_font(r, size=9.5, bold=True, color=ACCENT_DARK)

    for row_data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            p = cells[idx].paragraphs[0]
            p.style = doc.styles["Table Text"]
            r = p.add_run(value)
            set_run_font(r, size=9.5)
    set_table_geometry(table, widths_dxa)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_requirement(doc: Document, code: str, title: str, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.keep_together = True
    r = p.add_run(f"{code} {title}. ")
    set_run_font(r, bold=True, color=ACCENT_DARK)
    r = p.add_run(text)
    set_run_font(r)


def build_spec() -> None:
    doc = Document()
    configure_document(doc, "Simple WMS | Спецификация 0.1")
    add_cover(
        doc,
        kicker="Спецификация",
        title="Simple WMS 0.1",
        subtitle="Автономная система адресного складского управления",
        status="Рабочая спецификация Alpha",
        version="0.1 Alpha",
    )

    add_heading(doc, "1. Статус и назначение документа")
    add_paragraph(
        doc,
        "Настоящий документ определяет целевую модель, границы, обязательные функции и критерии готовности "
        "первой публичной версии Simple WMS. Он используется как основной источник требований к ядру продукта, "
        "интерфейсам, данным, развёртыванию и последующему модульному развитию."
    )
    add_callout(
        doc,
        "Simple WMS является самостоятельной складской системой. Для выполнения основных складских операций "
        "подключение ERP или иной учётной системы не требуется."
    )
    add_paragraph(
        doc,
        "Интеграции с 1С, другими ERP, корпоративными шинами и государственными системами прослеживаемости "
        "подключаются отдельными адаптерами и не должны изменять правила складского ядра."
    )

    add_heading(doc, "2. Видение продукта")
    add_paragraph(
        doc,
        "Simple WMS предназначена для предприятий, которым необходимо организовать адресный склад и "
        "прослеживаемость физических единиц без обязательного внедрения тяжёлой корпоративной платформы."
    )
    add_bullets(
        doc,
        [
            "автономная работа на обычном сервере под Linux с PostgreSQL;",
            "использование стандартных браузеров, сканеров, мобильных устройств и подключаемого оборудования;",
            "единое backend-ядро для web-интерфейса, ТСД, печати и интеграций;",
            "настраиваемые справочники вместо отраслевых констант в коде;",
            "полная история действий и запрет бесследных ручных исправлений;",
            "одно ядро продукта с подключаемыми функциональными и отраслевыми модулями.",
        ],
    )

    add_heading(doc, "3. Границы версии 0.1 Alpha")
    add_heading(doc, "3.1. В состав версии входят", 2)
    add_bullets(
        doc,
        [
            "справочники организаций, складов, зон, адресов, товаров, партий, единиц измерения и типов логистических единиц;",
            "приёмка, формирование логистических единиц, размещение, перемещение и межскладская передача;",
            "резервирование, экспедиция, контроль погрузки и завершение отгрузки;",
            "полная и циклическая инвентаризация без обязательной остановки склада;",
            "очередь складских заданий и компактное рабочее место оператора;",
            "карточки объектов, поиск, журнал операций, остатки и базовые отчёты;",
            "печатные формы и профили оборудования;",
            "пользователи, роли, права и ограничение доступа по складам;",
            "мастер первого запуска, демо-режим, установка, обновление и резервное копирование;",
            "REST API и базовый интеграционный контур для внешних адаптеров.",
        ],
    )
    add_heading(doc, "3.2. Не входят в обязательный объём", 2)
    add_bullets(
        doc,
        [
            "полноценный производственный, финансовый или бухгалтерский учёт;",
            "кадровый учёт, расчёт заработной платы и управление транспортом;",
            "гарантированная совместимость со всеми моделями ТСД, принтеров и весов;",
            "офлайн-синхронизация нескольких автономных серверов;",
            "готовый обмен с любой конфигурацией 1С без обследования и настройки;",
            "регламентированные операции «Честного ЗНАКА», ЕГАИС и иных государственных систем.",
        ],
    )

    add_heading(doc, "4. Архитектурные принципы")
    add_requirement(
        doc,
        "ARCH-01",
        "Единое ядро",
        "Все интерфейсы и адаптеры должны вызывать одинаковые backend-операции. "
        "Бизнес-запреты не должны существовать только в web-интерфейсе.",
    )
    add_requirement(
        doc,
        "ARCH-02",
        "Модульность",
        "Отраслевые функции, интеграции и драйверы оборудования подключаются поверх стабильных доменных интерфейсов.",
    )
    add_requirement(
        doc,
        "ARCH-03",
        "Транзакционность",
        "Изменение местоположения, состава, статуса и остатков выполняется атомарно с созданием события аудита.",
    )
    add_requirement(
        doc,
        "ARCH-04",
        "Расширяемые справочники",
        "Типы упаковки, единицы измерения, адресные уровни, префиксы и профили этикеток не задаются закрытыми списками.",
    )
    add_requirement(
        doc,
        "ARCH-05",
        "Идемпотентность",
        "Повторная команда сканера или интеграции не должна создавать дубликат физического объекта либо движения.",
    )

    add_heading(doc, "5. Основные пользователи и права")
    add_table(
        doc,
        ["Роль", "Основные возможности"],
        [
            ["Оператор приёмки", "Приёмка, проверка и формирование первичных логистических единиц."],
            ["Кладовщик", "Размещение, перемещение, отбор и выполнение назначенных заданий."],
            ["Оператор отгрузки", "Экспедиция, контроль погрузки и завершение отгрузки."],
            ["Старший кладовщик", "Исключительные операции, разукомплектация и обработка расхождений."],
            ["Руководитель склада", "Задания, инвентаризация, отчёты, разрешение блокировок и контроль сотрудников."],
            ["Администратор", "Пользователи, справочники, настройки, оборудование, интеграции и обслуживание."],
            ["Аудитор/наблюдатель", "Просмотр карточек, истории и отчётов без изменения складских данных."],
        ],
        [2200, CONTENT_WIDTH_DXA - 2200],
    )
    add_requirement(
        doc,
        "SEC-01",
        "Аутентификация",
        "Каждый пользователь должен входить под собственной учётной записью. Общая учётная запись склада не допускается.",
    )
    add_requirement(
        doc,
        "SEC-02",
        "Авторизация",
        "Права проверяются на backend для каждой изменяющей операции; скрытие кнопки не считается ограничением доступа.",
    )
    add_requirement(
        doc,
        "SEC-03",
        "Область доступа",
        "Пользователю назначаются доступные организации и склады. Межскладские действия требуют прав на обе стороны.",
    )

    add_heading(doc, "6. Справочники и модель данных")
    add_heading(doc, "6.1. Организация и настройки", 2)
    add_bullets(
        doc,
        [
            "наименование и реквизиты организации;",
            "часовой пояс, язык интерфейса и форматы даты/чисел;",
            "префиксы и правила нумерации объектов;",
            "единицы измерения и форматы этикеток по умолчанию;",
            "режим работы: демонстрационный или продуктивный.",
        ],
    )

    add_heading(doc, "6.2. Единицы измерения", 2)
    add_paragraph(
        doc,
        "Единица измерения является самостоятельным объектом справочника. Пользователь может добавлять "
        "собственные единицы, задавать обозначение, точность и принадлежность к группе измерений."
    )
    add_table(
        doc,
        ["Группа", "Примеры", "Правило"],
        [
            ["Количество", "шт., упаковка, комплект", "Допускаются целые и дробные значения согласно точности единицы."],
            ["Масса", "мг, г, кг, т", "Пересчёт выполняется через базовую единицу массы."],
            ["Объём", "мл, л, м³", "Пересчёт выполняется через базовую единицу объёма."],
            ["Длина/площадь", "м, м²", "Используются для соответствующих видов продукции."],
        ],
        [1700, 2500, CONTENT_WIDTH_DXA - 4200],
    )
    add_requirement(
        doc,
        "UOM-01",
        "Точность",
        "Количество и коэффициенты пересчёта хранятся как десятичные значения без применения float.",
    )
    add_requirement(
        doc,
        "UOM-02",
        "Пересчёт",
        "Для товара могут задаваться упаковочные коэффициенты, например 1 канистра = 20 л или 1 коробка = 12 шт.",
    )
    add_requirement(
        doc,
        "UOM-03",
        "Неизменность истории",
        "Изменение коэффициента не должно пересчитывать уже проведённые операции; в событии сохраняется применённое значение.",
    )

    add_heading(doc, "6.3. Типы логистических единиц", 2)
    add_paragraph(
        doc,
        "Логистическая единица представляет физически идентифицируемый носитель товара: коробку, ящик, "
        "палету, бочку, канистру, рулон, контейнер или еврокуб. Типы создаются и изменяются администратором."
    )
    add_bullets(
        doc,
        [
            "код, наименование и префикс идентификатора;",
            "масса тары, длина, ширина, высота и расчётный объём;",
            "допустимая масса и вместимость;",
            "признак возможности содержать товар либо другие логистические единицы;",
            "допустимые типы вложенных единиц;",
            "шаблон штрихкода и профиль этикетки;",
            "признак оборотной тары и статус активности.",
        ],
    )
    add_callout(
        doc,
        "Состав хранится как дерево вложенности: палета → коробки → товар; палета → бочки → литры; "
        "еврокуб → литры. Фиксированная связка «палета-коробка» считается частным случаем общей модели."
    )
    add_requirement(
        doc,
        "LU-01",
        "Уникальность",
        "Каждая отслеживаемая логистическая единица получает уникальный идентификатор в пределах системы.",
    )
    add_requirement(
        doc,
        "LU-02",
        "Циклы",
        "Система запрещает вложение единицы в саму себя и создание циклической структуры.",
    )
    add_requirement(
        doc,
        "LU-03",
        "Одна принадлежность",
        "Единица не может одновременно входить в несколько родительских единиц.",
    )
    add_requirement(
        doc,
        "LU-04",
        "Прослеживаемость",
        "Разукомплектация, объединение и переупаковка сохраняют связь между исходными и новыми объектами.",
    )

    add_heading(doc, "6.4. Товар и партия", 2)
    add_bullets(
        doc,
        [
            "код, наименование, основная и альтернативные единицы измерения;",
            "штрихкоды и внешние идентификаторы;",
            "правила упаковки и допустимые типы логистических единиц;",
            "срок годности и применяемая стратегия отбора;",
            "партия, дата производства, срок годности и статус качества;",
            "ограничения совместного хранения и дополнительные атрибуты.",
        ],
    )
    add_requirement(
        doc,
        "ITEM-01",
        "Источник данных",
        "Товары и партии могут создаваться вручную, импортироваться из CSV/XLSX либо поступать из внешней системы.",
    )
    add_requirement(
        doc,
        "ITEM-02",
        "Внешние коды",
        "Объект может иметь несколько внешних идентификаторов с указанием системы-источника, включая 1С и системы маркировки.",
    )

    add_heading(doc, "6.5. Адресная структура склада", 2)
    add_callout(
        doc,
        "Базовая иерархия: организация → склад → зона → проход → стеллаж → секция → ярус → позиция."
    )
    add_paragraph(
        doc,
        "Для напольного хранения и нестандартных помещений допускается пропуск отдельных уровней. "
        "Код адреса формируется по настраиваемому шаблону и не зависит от координат объекта на карте."
    )
    add_bullets(
        doc,
        [
            "тип местоположения: приёмка, хранение, карантин, расхождения, экспедиция, транзит, брак;",
            "вместимость по количеству мест, массе, объёму и габаритам;",
            "допустимые типы логистических единиц и категории товара;",
            "признак активности и временной блокировки;",
            "координаты и поворот на карте склада как отдельное представление адреса.",
        ],
    )

    add_heading(doc, "7. Функциональные процессы")
    add_heading(doc, "7.1. Первоначальная настройка", 2)
    add_numbers(
        doc,
        [
            "Создать первоначального администратора.",
            "Указать организацию, склад, часовой пояс и правила нумерации.",
            "Создать либо импортировать зоны и адреса.",
            "Заполнить единицы измерения и типы логистических единиц.",
            "Создать либо импортировать товары и партии.",
            "Проверить принтеры, сканеры и шаблоны этикеток.",
            "Зафиксировать завершение настройки и открыть рабочее место.",
        ],
    )

    add_heading(doc, "7.2. Приёмка и формирование единиц", 2)
    add_bullets(
        doc,
        [
            "создание или получение ожидаемой поставки/выпуска;",
            "идентификация товара, партии, количества и единицы измерения;",
            "создание логистической единицы и печать её этикетки;",
            "сканирование и подтверждение фактической приёмки;",
            "вложение единиц друг в друга с проверкой вместимости и совместимости;",
            "закрытие сформированной единицы и постановка задания на размещение.",
        ],
    )

    add_heading(doc, "7.3. Размещение и перемещение", 2)
    add_numbers(
        doc,
        [
            "Оператор сканирует перемещаемую единицу.",
            "Система показывает исходное место и рекомендует допустимое место назначения.",
            "Оператор сканирует адрес назначения.",
            "Backend проверяет склад, доступность, вместимость, ограничения и статус единицы.",
            "Перемещение фиксируется транзакционно, исходное место освобождается, задание закрывается.",
        ],
    )

    add_heading(doc, "7.4. Складские задания", 2)
    add_bullets(
        doc,
        [
            "типы: формирование, размещение, перемещение, отбор, отгрузка, инвентаризация, межскладская передача;",
            "приоритет, склад, основание, объект, исполнитель и контрольные сроки;",
            "состояния: новое, назначено, выполняется, приостановлено, завершено, отменено;",
            "автоматическое закрытие после успешной бизнес-операции;",
            "возврат рабочего места к очереди и предложение следующего приоритетного задания.",
        ],
    )

    add_heading(doc, "7.5. Резервирование и отгрузка", 2)
    add_bullets(
        doc,
        [
            "создание заявки вручную либо получение из внешней системы;",
            "автоматический или визуальный выбор доступных единиц;",
            "резервирование с учётом статуса качества и стратегии FIFO/FEFO;",
            "перемещение в экспедицию;",
            "контроль погрузки сканированием каждой назначенной единицы;",
            "завершение отгрузки после фактического окончания погрузки;",
            "освобождение складских адресов и фиксация отгруженного остатка.",
        ],
    )

    add_heading(doc, "7.6. Межскладская передача", 2)
    add_paragraph(
        doc,
        "Передача должна различать резервирование, подготовку, погрузку, нахождение в пути, приёмку "
        "складом назначения и последующее размещение. До подтверждения приёмки товар имеет статус «в пути»."
    )

    add_heading(doc, "7.7. Инвентаризация", 2)
    add_bullets(
        doc,
        [
            "полная, зональная и циклическая инвентаризация без обязательной остановки склада;",
            "последовательность «скан адреса → скан единицы» либо подтверждение пустого адреса;",
            "прогресс, непроверенные адреса и непроверенные единицы;",
            "расхождения: недостача, излишек, чужой адрес, неизвестная единица;",
            "запрет завершения полной инвентаризации при непроверенных обязательных адресах;",
            "оформляемые действия: подтвердить недостачу, разместить найденное, переместить по факту, отправить в карантин.",
        ],
    )

    add_heading(doc, "7.8. Карточки, поиск и история", 2)
    add_bullets(
        doc,
        [
            "карточка товара, партии, логистической единицы, адреса, задания, отгрузки и инвентаризации;",
            "текущее состояние, состав, местоположение, связанные документы и события;",
            "поиск по внутреннему коду, штрихкоду и внешнему идентификатору;",
            "неизменяемая история операций с пользователем, временем, причиной и снимками до/после;",
            "печать выбранной этикетки или документа непосредственно из карточки.",
        ],
    )

    add_heading(doc, "8. Остатки, отчёты и стратегии")
    add_heading(doc, "8.1. Базовые отчёты версии 0.1", 2)
    add_bullets(
        doc,
        [
            "остатки по организации, складу, зоне, адресу, товару, партии и статусу;",
            "физический, доступный, зарезервированный, заблокированный и транзитный остаток;",
            "движения за период с фильтрами по объекту и пользователю;",
            "сроки годности и партии, требующие внимания;",
            "незавершённые задания, неразмещённые единицы и проблемные строки инвентаризации;",
            "выгрузка выбранного отчёта в CSV/XLSX.",
        ],
    )
    add_heading(doc, "8.2. Стратегии", 2)
    add_requirement(
        doc,
        "STR-01",
        "FIFO",
        "Рекомендация к отбору по времени поступления либо создания остатка.",
    )
    add_requirement(
        doc,
        "STR-02",
        "FEFO",
        "Рекомендация к отбору партии с ближайшим допустимым сроком годности.",
    )
    add_requirement(
        doc,
        "STR-03",
        "Размещение",
        "Рекомендуемый адрес выбирается среди допустимых мест с учётом зоны, вместимости, совместимости и занятости.",
    )
    add_requirement(
        doc,
        "STR-04",
        "Совместимость",
        "Запрещается размещение несовместимых товаров или статусов в одном адресе при наличии соответствующего правила.",
    )

    add_heading(doc, "9. Оборудование и этикетки")
    add_heading(doc, "9.1. Общий принцип", 2)
    add_paragraph(
        doc,
        "Поддержка оборудования реализуется через профили и драйверы. Отсутствие конкретного драйвера "
        "не должно блокировать ручное выполнение операции при наличии соответствующего права."
    )
    add_table(
        doc,
        ["Класс", "Базовые способы подключения", "Объём версии 0.1"],
        [
            ["Принтеры", "PDF, системная очередь, RAW TCP", "Профили форматов, тестовая печать, АТОЛ ТТ42/HPRT XT100 как проверенный профиль."],
            ["Сканеры", "Клавиатурный ввод, камера, web/ТСД", "Стандартный ввод со сканера и компактный web-интерфейс."],
            ["Весы", "Ручной ввод, COM/USB, TCP", "Профиль устройства и ручной ввод; аппаратные драйверы подключаются отдельно."],
            ["ТСД", "Мобильный браузер, PWA, специализированное ПО", "Адаптивное рабочее место; нативные клиенты не обязательны."],
        ],
        [1500, 2800, CONTENT_WIDTH_DXA - 4300],
    )
    add_requirement(
        doc,
        "DEV-01",
        "Профиль",
        "Профиль содержит класс устройства, адрес/очередь, протокол, параметры подключения, склад и статус активности.",
    )
    add_requirement(
        doc,
        "DEV-02",
        "Диагностика",
        "Администратор может выполнить проверку доступности и тестовую печать без проведения складской операции.",
    )
    add_requirement(
        doc,
        "LBL-01",
        "Шаблоны",
        "Формат этикетки выбирается по типу единицы и профилю принтера; поддерживается печать выбранных объектов.",
    )
    add_requirement(
        doc,
        "LBL-02",
        "Содержимое кода",
        "По умолчанию машиночитаемый код содержит только устойчивый идентификатор объекта без URL и лишних данных.",
    )

    add_heading(doc, "10. Интеграционный контур")
    add_callout(
        doc,
        "Интеграция является подключаемой возможностью. Сбой внешней системы не должен повреждать складские данные "
        "или останавливать автономные операции, для которых внешний ответ не является обязательным."
    )
    add_bullets(
        doc,
        [
            "версионируемый REST API;",
            "webhooks для уведомления о складских событиях;",
            "журнал входящих и исходящих сообщений;",
            "надёжная очередь исходящих событий (outbox);",
            "идемпотентные ключи и защита от повторной обработки;",
            "повторная отправка, диагностика и ручной перезапуск сообщения;",
            "сопоставление внутренних и внешних идентификаторов.",
        ],
    )
    add_heading(doc, "10.1. Адаптер 1С", 2)
    add_paragraph(
        doc,
        "1С является приоритетным направлением интеграции для российского рынка и стран СНГ. Конкретный "
        "способ обмена определяется после обследования конфигурации и может использовать HTTP-сервисы, "
        "расширение, регламентную обработку либо файловый обмен."
    )
    add_bullets(
        doc,
        [
            "1С → WMS: товары, единицы, контрагенты, заказы, партии и иные мастер-данные;",
            "WMS → 1С: фактическая приёмка, отбор, отгрузка, межскладская передача и результаты инвентаризации;",
            "внутренние перемещения между ячейками по умолчанию остаются в WMS;",
            "прямая запись в таблицы базы 1С запрещается.",
        ],
    )
    add_heading(doc, "10.2. Регуляторная прослеживаемость", 2)
    add_paragraph(
        doc,
        "В перспективе предусматриваются коммерческие модули интеграции с «Честным ЗНАКом», ЕГАИС и "
        "иными государственными системами. Ядро должно хранить внешние коды маркировки, статусы и связь "
        "с логистическими единицами, но регламентные документы, криптография и обмен реализуются отдельными адаптерами."
    )

    add_heading(doc, "11. Демо-режим и первоначальные данные")
    add_bullets(
        doc,
        [
            "генерация управляемого набора складов, адресов, товаров, партий, единиц и операций;",
            "отдельная маркировка демонстрационных данных;",
            "восстановление исходной демонстрационной схемы с очисткой связанных объектов;",
            "запрет операций полной очистки в продуктивном режиме;",
            "явное отображение текущего режима в административном интерфейсе.",
        ],
    )

    add_heading(doc, "12. Развёртывание, обновление и эксплуатация")
    add_heading(doc, "12.1. Целевая конфигурация", 2)
    add_bullets(
        doc,
        [
            "Debian/Ubuntu либо совместимый Linux;",
            "Python в изолированном виртуальном окружении;",
            "PostgreSQL как основная СУБД;",
            "systemd для запуска приложения и резервного копирования;",
            "reverse proxy и HTTPS для продуктивного доступа;",
            "внешний каталог для резервных копий и журналов.",
        ],
    )
    add_heading(doc, "12.2. Автоматизированная установка", 2)
    add_numbers(
        doc,
        [
            "Проверить операционную систему, зависимости и свободные ресурсы.",
            "Создать системного пользователя, каталоги, виртуальное окружение и базу PostgreSQL.",
            "Сформировать конфигурацию и секреты с безопасными правами доступа.",
            "Применить миграции Alembic.",
            "Установить службы приложения и резервного копирования.",
            "Создать первоначального администратора.",
            "Проверить health endpoint и открыть мастер первого запуска.",
        ],
    )
    add_heading(doc, "12.3. Безопасное обновление", 2)
    add_callout(doc, "Рекомендуемая последовательность: backup → install → migrate → restart → health check → smoke test.")
    add_requirement(
        doc,
        "OPS-01",
        "Миграции",
        "Изменение структуры базы выполняется только версионированными миграциями.",
    )
    add_requirement(
        doc,
        "OPS-02",
        "Резервная копия",
        "Перед обновлением создаётся резервная копия; процедура восстановления документируется и периодически проверяется.",
    )
    add_requirement(
        doc,
        "OPS-03",
        "Версия",
        "Интерфейс и health endpoint показывают версию приложения и состояние миграций.",
    )

    add_heading(doc, "13. Нефункциональные требования")
    add_requirement(
        doc,
        "NFR-01",
        "Целостность",
        "Критические операции выполняются в транзакции и не оставляют частично изменённых объектов.",
    )
    add_requirement(
        doc,
        "NFR-02",
        "Аудит",
        "Изменяющее действие фиксирует пользователя, время, объект, основание и значения до/после.",
    )
    add_requirement(
        doc,
        "NFR-03",
        "Время",
        "События хранятся в UTC, а отображаются в часовом поясе склада или пользователя.",
    )
    add_requirement(
        doc,
        "NFR-04",
        "Совместимость",
        "Основной web-интерфейс поддерживает актуальные версии Chromium-браузеров на компьютере, телефоне и ТСД.",
    )
    add_requirement(
        doc,
        "NFR-05",
        "Наблюдаемость",
        "Доступны health endpoint, структурированный журнал ошибок и диагностика интеграций и оборудования.",
    )
    add_requirement(
        doc,
        "NFR-06",
        "Тестирование",
        "Доменные запреты и основные сквозные сценарии покрываются автоматическими тестами.",
    )
    add_requirement(
        doc,
        "NFR-07",
        "Конфиденциальность",
        "Пароли, ключи и строки подключения не хранятся в репозитории или журнале операций.",
    )

    add_heading(doc, "14. Критические бизнес-правила")
    add_bullets(
        doc,
        [
            "один идентификатор соответствует одному физическому объекту;",
            "одна логистическая единица имеет не более одного непосредственного родителя и одного текущего местоположения;",
            "состав закрытой единицы меняется только исключительной операцией с указанием причины;",
            "заблокированный, карантинный, списанный или отгруженный товар нельзя использовать в обычном отборе;",
            "зарезервированный товар нельзя назначить другому активному документу;",
            "перемещение в занятое либо несовместимое место запрещается;",
            "завершённая отгрузка освобождает адреса и исключает единицы из складского остатка;",
            "ручная корректировка не изменяет прошлые события и всегда создаёт новое событие аудита;",
            "повторное сканирование не создаёт дубликат операции.",
        ],
    )

    add_heading(doc, "15. Критерии готовности версии 0.1 Alpha")
    add_table(
        doc,
        ["Область", "Критерий приёмки"],
        [
            ["Установка", "Чистый поддерживаемый сервер разворачивается по документированной процедуре до рабочего входа в систему."],
            ["Настройка", "Администратор создаёт организацию, склад, адреса, единицы измерения и типы логистических единиц через UI."],
            ["Права", "Оператор не может выполнить административную или исключительную операцию ни через UI, ни через API."],
            ["Складской цикл", "Приёмка → формирование → размещение → перемещение → отбор → погрузка выполняются с полной историей."],
            ["Инвентаризация", "Пустые, совпавшие и проблемные адреса обрабатываются; незавершённый обход виден пользователю."],
            ["Отчёты", "Остатки и движения доступны по складу, адресу, товару, партии и статусу."],
            ["Оборудование", "Сканер-клавиатура и проверенный профиль термопринтера проходят сквозной тест."],
            ["Эксплуатация", "Резервная копия создаётся автоматически, а тестовое восстановление выполняется по инструкции."],
            ["Качество", "Миграции находятся на head, автоматические тесты проходят, health endpoint доступен локально и по LAN."],
        ],
        [2100, CONTENT_WIDTH_DXA - 2100],
    )

    add_heading(doc, "16. Направления последующего развития")
    add_bullets(
        doc,
        [
            "расширенные стратегии размещения, пополнения и отбора;",
            "контроль качества, карантин, сертификаты и фотографии партии;",
            "контроль сроков годности и отраслевые правила FEFO;",
            "нативные клиенты ТСД и ограниченный офлайн-режим;",
            "драйверы весов, промышленных принтеров и стационарных сканеров;",
            "конструктор отчётов и показателей эффективности склада;",
            "адаптеры 1С для согласованных конфигураций;",
            "модули «Честный ЗНАК», ЕГАИС и иной регуляторной прослеживаемости;",
            "исключительно коммерческий кассовый модуль (POS) с рабочим местом кассира, сменами, продажами, возвратами и фискальным оборудованием;",
            "отраслевые пакеты для химии, пищевого производства и дистрибуции.",
        ],
    )
    add_callout(
        doc,
        "Продукт развивается как одно ядро с подключаемыми возможностями. Создание самостоятельных расходящихся "
        "версий для отдельных заказчиков не является целевой архитектурой."
    )

    add_heading(doc, "17. Открытые продуктовые решения")
    add_bullets(
        doc,
        [
            "лицензия и модель распространения публичной версии;",
            "минимальные поддерживаемые характеристики сервера и целевая нагрузка;",
            "границы офлайн-работы ТСД;",
            "политика совместимости драйверов оборудования;",
            "правила версионирования API и модулей;",
            "формат поставки: установочный сценарий, пакет, контейнер либо несколько вариантов.",
        ],
    )

    SPEC_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(SPEC_PATH)


TERMS_GROUPS: list[tuple[str, list[tuple[str, str]]]] = [
    (
        "1. Продукт и архитектура",
        [
            ("Simple WMS", "Автономная система управления складскими операциями и адресным хранением."),
            ("WMS", "Система управления складом: размещением, движениями, заданиями, запасами и прослеживаемостью."),
            ("Автономная работа", "Выполнение основных складских операций без обязательного подключения ERP."),
            ("ERP", "Внешняя система управления ресурсами предприятия, учётом и документами."),
            ("Складское ядро", "Набор доменных правил и операций, общий для UI, ТСД, оборудования и интеграций."),
            ("Модуль", "Подключаемая возможность, расширяющая ядро без изменения его базовых правил."),
            ("Адаптер", "Компонент связи с конкретной внешней системой, протоколом или оборудованием."),
            ("Версия Alpha", "Ранняя рабочая версия для контролируемой эксплуатации и обратной связи; интерфейсы могут изменяться."),
        ],
    ),
    (
        "2. Количество, товар и упаковка",
        [
            ("Товар", "Номенклатурная позиция, для которой ведутся складские остатки и движения."),
            ("Партия", "Группа товара с общими признаками происхождения, производства, срока годности или качества."),
            ("Единица измерения", "Правило количественного выражения товара: штука, килограмм, литр, метр и другое."),
            ("Базовая единица", "Единица, через которую выполняются пересчёты внутри одной группы измерений."),
            ("Коэффициент пересчёта", "Зафиксированное отношение одной единицы или упаковки к другой, например 1 канистра = 20 л."),
            ("Точность количества", "Допустимое число знаков после запятой для конкретной единицы измерения."),
            ("Тип логистической единицы", "Настраиваемый вид физического носителя: коробка, палета, ящик, бочка, канистра, рулон, контейнер или еврокуб."),
            ("Логистическая единица", "Идентифицируемый физический объект, который содержит товар либо другие логистические единицы."),
            ("Вложенность", "Связь родительской и дочерней единицы, например палета → коробки → товар."),
            ("Тара", "Физическая упаковка или ёмкость; может быть одноразовой либо оборотной."),
            ("Масса тары", "Масса пустой упаковки, используемая для расчёта нетто и брутто."),
            ("Вместимость", "Ограничение единицы или адреса по количеству мест, массе, объёму либо габаритам."),
            ("Разукомплектация", "Контролируемое изменение состава ранее закрытой логистической единицы."),
            ("Переупаковка", "Перенос товара в новую упаковку с сохранением связи с исходным объектом."),
        ],
    ),
    (
        "3. Адресное хранение",
        [
            ("Организация", "Владелец настроек, складов, пользователей и данных внутри одной установки."),
            ("Склад", "Территориально или организационно выделенное место хранения."),
            ("Зона", "Часть склада с назначением: приёмка, хранение, карантин, экспедиция и другое."),
            ("Проход", "Адресный уровень, объединяющий стеллажи или напольные позиции вдоль маршрута."),
            ("Стеллаж", "Конструкция для размещения секций, ярусов и позиций хранения."),
            ("Секция", "Вертикальная или логическая часть стеллажа."),
            ("Ярус", "Уровень хранения по высоте."),
            ("Позиция", "Минимальное адресуемое место размещения логистической единицы."),
            ("Адрес", "Уникальный код физического или логического местоположения на складе."),
            ("Напольное хранение", "Размещение без стеллажной иерархии в размеченных позициях пола."),
            ("Карта склада", "Визуальное представление зон, стеллажей и адресов; не является источником их идентичности."),
            ("Занятость", "Текущее использование вместимости адреса размещёнными единицами."),
        ],
    ),
    (
        "4. Складские процессы",
        [
            ("Приёмка", "Подтверждение фактического поступления товара или логистической единицы под ответственность склада."),
            ("Формирование", "Создание состава логистической единицы путём вложения товара или дочерних единиц."),
            ("Закрытие", "Фиксация состава единицы и перевод её в состояние, готовое к размещению или дальнейшей операции."),
            ("Размещение", "Перемещение принятой единицы из зоны приёмки в адрес хранения."),
            ("Перемещение", "Контролируемая смена текущего адреса единицы."),
            ("Резервирование", "Закрепление доступного остатка за конкретной заявкой, отгрузкой или передачей."),
            ("Отбор", "Изъятие выбранных единиц из хранения для выполнения документа."),
            ("Экспедиция", "Зона и этап, в котором отобранные единицы ожидают контроля и погрузки."),
            ("Погрузка", "Фактическое помещение единиц в транспорт с контролем принадлежности документу."),
            ("Отгрузка", "Завершённая складская операция после фактического окончания погрузки."),
            ("Межскладская передача", "Отправка единиц со склада-источника и приёмка складом назначения."),
            ("Товар в пути", "Остаток, отправленный складом-источником, но ещё не принятый складом назначения."),
            ("Списание", "Оформленный вывод товара из доступного и физического складского остатка."),
        ],
    ),
    (
        "5. Задания и управление работой",
        [
            ("Складское задание", "Поручение выполнить определённую физическую операцию с объектом склада."),
            ("Очередь заданий", "Упорядоченный перечень доступных и назначенных работ."),
            ("Приоритет", "Степень срочности задания: низкая, обычная, высокая или срочная."),
            ("Исполнитель", "Пользователь, назначенный на выполнение задания."),
            ("Исключительная операция", "Действие вне обычного процесса, требующее специальных прав и указания причины."),
            ("Автоматическое закрытие", "Завершение задания после успешного выполнения связанной бизнес-операции."),
            ("Рабочее место", "Интерфейс оператора с очередью, текущим заданием и пошаговым сценарием."),
            ("ТСД", "Терминал сбора данных либо иное мобильное устройство, используемое для складских операций."),
        ],
    ),
    (
        "6. Остатки, статусы и стратегии",
        [
            ("Физический остаток", "Весь товар, который фактически находится на складе, включая резерв и блокировки."),
            ("Доступный остаток", "Товар, разрешённый для нового резервирования и отбора."),
            ("Зарезервированный остаток", "Товар, уже закреплённый за активным документом."),
            ("Транзитный остаток", "Товар в процессе межскладской передачи."),
            ("Карантин", "Временный запрет обычных операций до проверки или решения ответственного пользователя."),
            ("Блокировка", "Запрет определённых операций с товаром, партией, единицей или адресом."),
            ("Статус качества", "Состояние допуска партии или остатка к складским операциям и отгрузке."),
            ("FIFO", "Стратегия отбора в порядке поступления или создания остатка."),
            ("FEFO", "Стратегия отбора по ближайшему допустимому сроку годности."),
            ("Стратегия размещения", "Правила выбора рекомендуемого адреса с учётом вместимости и ограничений."),
            ("Совместимость хранения", "Правило, разрешающее или запрещающее совместное размещение определённых товаров или статусов."),
        ],
    ),
    (
        "7. Инвентаризация и расхождения",
        [
            ("Инвентаризация", "Сравнение ожидаемого и фактического состояния склада."),
            ("Полная инвентаризация", "Обход всех обязательных адресов выбранного склада."),
            ("Циклическая инвентаризация", "Регулярная проверка части склада без полной остановки операций."),
            ("Проверенный адрес", "Адрес, для которого подтверждён фактический состав либо пустое состояние."),
            ("Непроверенный адрес", "Адрес, включённый в инвентаризацию, но ещё не подтверждённый оператором."),
            ("Расхождение", "Несоответствие ожидаемого и фактического местоположения или количества."),
            ("Недостача", "Ожидаемый объект или количество не обнаружены."),
            ("Излишек", "Обнаружен объект или количество, отсутствующие в ожидаемом составе."),
            ("Чужой адрес", "Единица обнаружена не в том месте, которое указано в системе."),
            ("Корректирующая операция", "Оформленное действие по результатам подтверждённого расхождения."),
        ],
    ),
    (
        "8. Идентификация и оборудование",
        [
            ("Штрихкод", "Машиночитаемое графическое представление идентификатора или данных объекта."),
            ("Code 128", "Одномерная символика, подходящая для печати идентификаторов."),
            ("QR-код", "Двумерная символика, используемая для компактного кодирования идентификатора."),
            ("DataMatrix", "Компактная двумерная символика промышленного применения."),
            ("Этикетка", "Печатная форма с человекочитаемой и машиночитаемой информацией."),
            ("Профиль этикетки", "Настройка размера, состава и шаблона этикетки для типа объекта и принтера."),
            ("Профиль устройства", "Настройка класса, адреса, протокола и параметров оборудования."),
            ("RAW TCP", "Передача готовой команды печати напрямую на сетевой порт принтера."),
            ("Сканер-клавиатура", "Сканер, передающий считанный код как последовательность клавиш."),
            ("Системная очередь печати", "Служба операционной системы, управляющая заданиями принтера."),
        ],
    ),
    (
        "9. Пользователи, аудит и безопасность",
        [
            ("Аутентификация", "Проверка личности пользователя при входе."),
            ("Авторизация", "Проверка права пользователя выполнить конкретное действие."),
            ("Роль", "Набор разрешений, соответствующий рабочим обязанностям."),
            ("Область доступа", "Организации и склады, с которыми разрешено работать пользователю."),
            ("Журнал операций", "Последовательность событий об изменениях объектов и выполненных действиях."),
            ("История объекта", "События, относящиеся к конкретной единице, партии, адресу, заданию или документу."),
            ("Снимок до/после", "Зафиксированные значения объекта перед операцией и после неё."),
            ("Ручная корректировка", "Изменение данных уполномоченным пользователем с обязательной причиной и аудитом."),
        ],
    ),
    (
        "10. Интеграции",
        [
            ("Интеграционный контур", "API, очередь, журнал и правила обмена, общие для внешних адаптеров."),
            ("REST API", "Версионируемый программный интерфейс на основе HTTP."),
            ("Webhook", "Исходящий HTTP-вызов при наступлении складского события."),
            ("Outbox", "Транзакционный журнал событий, ожидающих надёжной отправки во внешнюю систему."),
            ("Идемпотентность", "Повторная обработка одной команды не создаёт повторный объект или движение."),
            ("Внешний идентификатор", "Код объекта в другой системе с указанием системы-источника."),
            ("Адаптер 1С", "Подключаемый компонент обмена между Simple WMS и согласованной конфигурацией 1С."),
            ("Кассовый модуль (POS)", "Исключительно коммерческое подключаемое рабочее место кассира для смен, продаж, возвратов, оплат и фискального оборудования."),
            ("«Честный ЗНАК»", "Внешняя государственная система маркировки; интеграция реализуется отдельным модулем."),
            ("ЕГАИС", "Внешняя государственная информационная система; интеграция реализуется отдельным модулем."),
            ("Журнал интеграции", "История входящих и исходящих сообщений, подтверждений, повторов и ошибок."),
        ],
    ),
    (
        "11. Эксплуатация",
        [
            ("Миграция базы", "Версионированное изменение структуры данных при установке или обновлении."),
            ("Резервная копия", "Копия данных и настроек, пригодная для восстановления."),
            ("Проверка восстановления", "Контрольная процедура, доказывающая пригодность резервной копии."),
            ("Health endpoint", "Технический адрес для проверки состояния приложения и его зависимостей."),
            ("Демо-режим", "Режим с управляемыми демонстрационными данными и разрешённым сбросом схемы."),
            ("Продуктивный режим", "Режим реальной эксплуатации, в котором опасные демонстрационные операции запрещены."),
            ("Мастер первого запуска", "Пошаговая настройка администратора, организации, склада и базовых справочников."),
            ("Smoke test", "Короткая проверка ключевого пользовательского сценария после установки или обновления."),
        ],
    ),
]


def build_terms() -> None:
    doc = Document()
    configure_document(doc, "Simple WMS | Термины и определения")
    add_cover(
        doc,
        kicker="Справочный документ",
        title="Термины и определения",
        subtitle="Единый словарь продукта Simple WMS",
        status="Приложение к спецификации Simple WMS 0.1",
        version="0.1 Alpha",
    )

    add_heading(doc, "Назначение словаря")
    add_paragraph(
        doc,
        "Документ устанавливает единое значение терминов, используемых в спецификации, интерфейсе, "
        "API, документации и обсуждении доработок Simple WMS. Термины конкретного заказчика могут "
        "добавляться отдельным приложением без изменения общих понятий продукта."
    )

    for heading, rows in TERMS_GROUPS:
        add_heading(doc, heading)
        add_table(
            doc,
            ["Термин", "Определение"],
            [[term, definition] for term, definition in rows],
            [2500, CONTENT_WIDTH_DXA - 2500],
        )

    TERMS_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(TERMS_PATH)


if __name__ == "__main__":
    build_spec()
    build_terms()
    print(SPEC_PATH)
    print(TERMS_PATH)
